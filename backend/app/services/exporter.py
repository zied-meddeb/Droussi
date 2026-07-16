import io
from xml.sax.saxutils import escape

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    ListFlowable,
    ListItem,
)

from ..models.schemas import ExamContent


def _pdf_text(value: str, *, keep_linebreaks: bool = False) -> str:
    """Escape user/LLM-provided text for ReportLab's mini-markup.

    ReportLab ``Paragraph`` interprets a subset of HTML-like tags, so raw ``<``,
    ``>``, and ``&`` from exam content could break rendering or inject markup.
    Escape first, then (optionally) re-introduce ``<br/>`` for real newlines.
    """
    escaped = escape(value or "")
    if keep_linebreaks:
        escaped = escaped.replace("\n", "<br/>")
    return escaped


def to_docx(exam: ExamContent) -> bytes:
    doc = Document()

    title = doc.add_heading(exam.title or "Exam", level=1)
    title.alignment = 1  # center

    p = doc.add_paragraph()
    p.alignment = 1
    run = p.add_run(f"Total: {exam.total_points} points")
    run.bold = True

    for i, ex in enumerate(exam.exercises, start=1):
        h = doc.add_heading(
            f"Exercise {i} ({'MCQ' if ex.type == 'mcq' else 'Open'}) — {ex.points} pts",
            level=2,
        )
        for r in h.runs:
            r.font.size = Pt(13)
        doc.add_paragraph(ex.question)
        if ex.type == "mcq" and ex.choices:
            for c in ex.choices:
                doc.add_paragraph(c, style="List Bullet")

    doc.add_page_break()
    doc.add_heading("Answer Key", level=1)
    for i, ex in enumerate(exam.exercises, start=1):
        para = doc.add_paragraph()
        para.add_run(f"Exercise {i}: ").bold = True
        para.add_run(ex.answer)
        if ex.explanation:
            ep = doc.add_paragraph()
            ep.add_run("Explanation: ").italic = True
            ep.add_run(ex.explanation)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def to_pdf(exam: ExamContent) -> bytes:
    buf = io.BytesIO()
    pdf = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title=exam.title or "Exam",
    )
    styles = getSampleStyleSheet()
    h1 = styles["Title"]
    h2 = ParagraphStyle("ExerciseHeader", parent=styles["Heading2"], spaceAfter=6)
    body = styles["BodyText"]

    story: list = [
        Paragraph(_pdf_text(exam.title or "Exam"), h1),
        Paragraph(f"<b>Total: {exam.total_points} points</b>", body),
        Spacer(1, 0.5 * cm),
    ]
    for i, ex in enumerate(exam.exercises, start=1):
        story.append(
            Paragraph(
                f"Exercise {i} ({'MCQ' if ex.type == 'mcq' else 'Open'}) "
                f"— {ex.points} pts",
                h2,
            )
        )
        story.append(Paragraph(_pdf_text(ex.question, keep_linebreaks=True), body))
        if ex.type == "mcq" and ex.choices:
            story.append(
                ListFlowable(
                    [ListItem(Paragraph(_pdf_text(c), body)) for c in ex.choices],
                    bulletType="bullet",
                )
            )
        story.append(Spacer(1, 0.3 * cm))

    story.append(Spacer(1, 0.5 * cm))
    story.append(Paragraph("Answer Key", h1))
    for i, ex in enumerate(exam.exercises, start=1):
        story.append(
            Paragraph(f"<b>Exercise {i}:</b> {_pdf_text(ex.answer)}", body)
        )
        if ex.explanation:
            story.append(
                Paragraph(
                    f"<i>Explanation:</i> {_pdf_text(ex.explanation)}", body
                )
            )
        story.append(Spacer(1, 0.2 * cm))

    pdf.build(story)
    return buf.getvalue()
